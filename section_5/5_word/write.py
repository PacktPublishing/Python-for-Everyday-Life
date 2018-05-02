# -*- coding: utf-8 -*-
# !/usr/bin/env python3

import pathlib
import docx


if __name__ == '__main__':
    composers_data = [
        {'surname': 'liszt', 'country': 'HU', 'name': 'franz', 'email': 'franz.liszt@gmail.com'},
        {'surname': 'bach', 'country': 'DE', 'name': 'johann sebastian', 'email': 'jsbach@yahoo.de'},
        {'surname': 'verdi', 'country': 'IT', 'name': 'giuseppe', 'email': 'giuseppeverdi@hotmail.com'},
        {'surname': 'gilmour', 'country': 'GB', 'name': 'david', 'email': 'dgilmour@me.com'}
    ]

    # Instantiating a Docx object
    doc = docx.Document()

    # Main header
    doc.add_heading('Composers Book', 0)

    # Adding composers data as bulleted lists
    doc.add_heading('List Form', 1)

    for composer in composers_data:
        doc.add_heading(composer['name'] + ' ' + composer['surname'], 2)
        doc.add_paragraph(composer['name'], style='ListBullet')
        doc.add_paragraph(composer['surname'], style='ListBullet')
        doc.add_paragraph(composer['country'], style='ListBullet')
        doc.add_paragraph(composer['email'], style='ListBullet')

    # Tables are also very useful!
    doc.add_heading('Tabular Form', 1)
    p = doc.add_paragraph('The following ')
    p.add_run('table ').bold = True
    p.add_run('contains ')
    p.add_run('composers ').italic = True
    p.add_run(' data')

    table = doc.add_table(rows=1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Surname'
    header_cells[2].text = 'Country'
    header_cells[3].text = 'Email'
    for composer in composers_data:
        row_cells = table.add_row().cells
        row_cells[0].text = composer['name']
        row_cells[1].text = composer['surname']
        row_cells[2].text = composer['country']
        row_cells[3].text = composer['email']

    # Saving our docx
    target = pathlib.Path('output.docx')
    target.touch()
    doc.save(str(target.resolve()))