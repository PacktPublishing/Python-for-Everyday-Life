# -*- coding: utf-8 -*-
# !/usr/bin/env python3

import pathlib
import docx
import pprint


def read_as_plain_text(docx_object):
    lines = [p.text for p in docx_object.paragraphs]
    return '\n'.join(lines)


if __name__ == '__main__':
    src = pathlib.Path('composers.docx').resolve()
    doc = docx.Document(str(src))

    # Print out the whole doc content as a string
    print('Whole Word doc content:\n')
    print(read_as_plain_text(doc))

    # Now let's turn the bullet lists into dictionaries
    print('\nReading composers info as dicts:\n')
    composers = []
    paragraphs = (p for p in doc.paragraphs)  # generator expression
    for p in paragraphs:
        if p.style.name == 'Heading 2':
            composer = dict()
            tokens = p.text.split(' ')
            composer['name'] = tokens[0].capitalize()
            composer['surname'] = tokens[1].capitalize()
            bullet_point_1 = next(paragraphs)
            composer['country'] = bullet_point_1.text.split(' ')[1]
            bullet_point_2 = next(paragraphs)
            composer['email'] = bullet_point_2.text.split(' ')[1]
            composers.append(composer)
    pprint.pprint(composers)